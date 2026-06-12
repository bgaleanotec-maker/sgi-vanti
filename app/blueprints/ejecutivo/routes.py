import io
import zipfile
from datetime import datetime
from flask import render_template, request, redirect, url_for, flash, send_file, abort
from flask_login import login_required, current_user
from docx import Document

from app.extensions import db
from app.models import Usuario, Imposibilidad
from app.decorators import role_required
from app.helpers import aplicar_filtros_comunes, guardar_datos_carta, guardar_soporte
from app.services.notification_service import notify_user
from app.blueprints.ejecutivo import ejecutivo_bp


@ejecutivo_bp.route('/')
@login_required
@role_required('ejecutivo')
def dashboard():
    """Ejecutivo ve TODAS las tareas donde es responsable:
    - tareas tipo 'carta' (para gestionar la carta final)
    - Y tambien tiene visibilidad de OTRAS tareas del mismo BP_Firma para ver el estado
      de los negocios de sus contratistas/firmas (pendiente, devuelto, gestionado, etc.)
    """
    from app.models.catalog import EstadoTareaConfig
    from sqlalchemy import or_

    # Tareas donde este ejecutivo es el asignado
    tareas_propias = Imposibilidad.query.filter_by(
        ejecutivo_asignado=current_user.username
    )

    # BPs_Firma vinculados a este ejecutivo (para ver TODAS las tareas de esas firmas)
    bps_del_ejecutivo = db.session.query(Imposibilidad.bp_firma).filter_by(
        ejecutivo_asignado=current_user.username
    ).distinct().all()
    bps = [b[0] for b in bps_del_ejecutivo if b[0]]

    # Tareas asociadas (mismas firmas del ejecutivo, para visibilidad de estado)
    if bps:
        tareas_asociadas_q = Imposibilidad.query.filter(Imposibilidad.bp_firma.in_(bps))
    else:
        tareas_asociadas_q = tareas_propias

    tareas_asociadas_q, filtros = aplicar_filtros_comunes(tareas_asociadas_q)

    # Filtros especificos del ejecutivo
    estado_filter = request.args.get('estado')
    bp_filter = request.args.get('bp_firma')
    vista = request.args.get('vista', 'todas')  # 'todas', 'cartas', 'propias'

    if vista == 'cartas':
        tareas_asociadas_q = tareas_asociadas_q.filter_by(tipo_tarea='carta')
    elif vista == 'propias':
        tareas_asociadas_q = tareas_asociadas_q.filter_by(ejecutivo_asignado=current_user.username)

    if estado_filter:
        tareas_asociadas_q = tareas_asociadas_q.filter_by(estado_tarea=estado_filter)
        filtros['estado'] = estado_filter
    if bp_filter:
        tareas_asociadas_q = tareas_asociadas_q.filter(
            Imposibilidad.bp_firma.ilike(f'%{bp_filter}%')
        )
        filtros['bp_firma'] = bp_filter

    tareas = tareas_asociadas_q.order_by(Imposibilidad.fecha_cargue.desc()).all()

    # Stats: conteo por estado para todas las tareas de los BPs vinculados
    estados = EstadoTareaConfig.query.filter_by(is_active=True).order_by(EstadoTareaConfig.order_index).all()
    # estado_map incluye TODOS los estados (activos + legacy) para badges historicos
    estado_map = {e.name: e for e in EstadoTareaConfig.query.all()}
    stats = {e.name: 0 for e in estados}
    total_cartas_pendientes = 0
    bps_summary = {}
    for t in tareas:
        if t.estado_tarea in stats:
            stats[t.estado_tarea] += 1
        if t.tipo_tarea == 'carta' and t.estado_tarea != 'carta_enviada':
            total_cartas_pendientes += 1
        bp = t.bp_firma or 'Sin BP'
        if bp not in bps_summary:
            bps_summary[bp] = {'total': 0, 'pendientes': 0, 'gestionados': 0, 'devueltas': 0}
        bps_summary[bp]['total'] += 1
        if t.estado_tarea in ('pendiente', 'recibida', 'escalado'):
            bps_summary[bp]['pendientes'] += 1
        elif t.estado_tarea in ('soportes_cargados', 'gestionado', 'finalizado', 'cerrada'):
            bps_summary[bp]['gestionados'] += 1
        elif t.estado_tarea in ('devuelta', 'rechazado', 'rechazada'):
            bps_summary[bp]['devueltas'] += 1

    return render_template(
        'dashboard_ejecutivo.html',
        tareas=tareas, user=current_user, filtros=filtros,
        estados=estados, estado_map=estado_map, stats=stats,
        total_cartas_pendientes=total_cartas_pendientes,
        bps_summary=bps_summary,
        vista=vista, estado_filter=estado_filter, bp_filter=bp_filter,
    )


@ejecutivo_bp.route('/tarea/<int:id>/solicitar_correccion', methods=['POST'])
@login_required
@role_required('ejecutivo')
def solicitar_correccion(id):
    """El ejecutivo interactua con el caso: reporta inconsistencias en los soportes
    o solicita informacion adicional a la firma. Adjunta archivo opcional, registra
    el comentario y regresa el caso a 'pendiente' notificando a la firma."""
    import os
    from app.config import Config
    tarea = Imposibilidad.query.get_or_404(id)
    # El ejecutivo solo puede actuar sobre tareas de los BPs que gestiona
    bps_del_ejecutivo = [b[0] for b in db.session.query(Imposibilidad.bp_firma).filter_by(
        ejecutivo_asignado=current_user.username).distinct().all() if b[0]]
    if tarea.bp_firma not in bps_del_ejecutivo and tarea.ejecutivo_asignado != current_user.username:
        abort(403)

    comentario = (request.form.get('comentario') or '').strip()
    if not comentario:
        flash('Debes indicar la inconsistencia o la informacion requerida.', 'warning')
        return redirect(url_for('ejecutivo.dashboard'))

    archivo = request.files.get('archivo')
    if archivo and archivo.filename:
        filename = f"{tarea.id}_ejec_{archivo.filename}"
        guardar_soporte(archivo, filename)
        tarea.archivo_nombre = filename

    sello = f"\n[EJECUTIVO {current_user.username} {datetime.now():%Y-%m-%d %H:%M}]: {comentario}"
    tarea.comentarios_gestor = (tarea.comentarios_gestor or '') + sello
    # Regresa el caso para que la firma corrija / cargue soportes claros
    tarea.estado_tarea = 'pendiente'
    db.session.commit()

    # Notificar a la firma / contratista
    from sqlalchemy import or_ as _or
    recipientes = Usuario.query.filter(
        _or(Usuario.username == tarea.bp_firma, Usuario.bp_firma == tarea.bp_firma),
        Usuario.is_active == True
    ).all()
    mensaje = (
        f"El ejecutivo solicita correccion en la orden {tarea.orden}. "
        f"Detalle: {comentario}. Ingresa a SGI Vanti para ajustar la informacion."
    )
    subject = f"SGI Vanti - Correccion requerida en orden {tarea.orden}"
    html = render_template('email_base.html', content=f"""
        <p>Hola,</p>
        <p>El ejecutivo ha revisado la orden <strong>{tarea.orden}</strong> y requiere ajustes:</p>
        <p style="background:#fff7ed;padding:10px;border-radius:6px;">{comentario}</p>
        <p>Por favor ingresa a la plataforma SGI Vanti para corregir o aclarar la informacion.</p>
    """)
    for u in recipientes:
        try:
            notify_user(u, subject, html, mensaje, tarea.id)
        except Exception as e:
            print(f"[ejecutivo.solicitar_correccion] error notificando {u.username}: {e}")

    flash(f'Solicitud de correccion enviada a la firma para la orden {tarea.orden}.', 'success')
    return redirect(url_for('ejecutivo.dashboard'))


@ejecutivo_bp.route('/tarea/<int:id>/cambiar_estado', methods=['POST'])
@login_required
@role_required('ejecutivo')
def cambiar_estado(id):
    """El ejecutivo cambia el estado del negocio (rechazado, escalado, anulado,
    finalizado, etc.), tipicamente cuando la firma ya cargo los soportes. Solo
    puede actuar sobre tareas de los BPs que gestiona. Notifica a la firma."""
    from app.models.catalog import EstadoTareaConfig
    tarea = Imposibilidad.query.get_or_404(id)

    # Scope: solo BPs del ejecutivo (o tareas asignadas a el)
    bps_del_ejecutivo = [b[0] for b in db.session.query(Imposibilidad.bp_firma).filter_by(
        ejecutivo_asignado=current_user.username).distinct().all() if b[0]]
    if tarea.bp_firma not in bps_del_ejecutivo and tarea.ejecutivo_asignado != current_user.username:
        abort(403)

    nuevo_estado = (request.form.get('nuevo_estado') or '').strip()
    comentario = (request.form.get('comentario') or '').strip()

    # Solo estados activos (dinamico/configurable desde catalogos)
    estados_validos = {e.name for e in EstadoTareaConfig.query.filter_by(is_active=True).all()}
    if nuevo_estado not in estados_validos:
        flash('Estado no valido.', 'warning')
        return redirect(url_for('ejecutivo.dashboard'))

    anterior = tarea.estado_tarea
    tarea.estado_tarea = nuevo_estado

    # Ajustes coherentes con el estado destino
    if nuevo_estado == 'rechazado':
        tarea.tipo_negacion = 'rechazo'
        if not tarea.clasificacion:
            tarea.clasificacion = 'INSO'
        if comentario:
            tarea.motivo_rechazo = comentario
    if nuevo_estado in ('rechazado', 'finalizado', 'anulado'):
        tarea.fecha_gestion_gestor = datetime.now()

    sello = f"\n[EJECUTIVO {current_user.username} {datetime.now():%Y-%m-%d %H:%M}] estado {anterior}->{nuevo_estado}"
    if comentario:
        sello += f": {comentario}"
    tarea.comentarios_gestor = (tarea.comentarios_gestor or '') + sello
    db.session.commit()

    # Notificar a la firma / contratista
    from sqlalchemy import or_ as _or
    recipientes = Usuario.query.filter(
        _or(Usuario.username == tarea.bp_firma, Usuario.bp_firma == tarea.bp_firma),
        Usuario.is_active == True
    ).all()
    estado_cfg = EstadoTareaConfig.query.filter_by(name=nuevo_estado).first()
    estado_label = estado_cfg.display_name if estado_cfg else nuevo_estado
    mensaje = (
        f"El ejecutivo actualizo la orden {tarea.orden} a estado '{estado_label}'."
        + (f" Nota: {comentario}." if comentario else "")
    )
    subject = f"SGI Vanti - Orden {tarea.orden} actualizada a {estado_label}"
    html = render_template('email_base.html', content=f"""
        <p>Hola,</p>
        <p>El ejecutivo actualizo la orden <strong>{tarea.orden}</strong> al estado
        <strong>{estado_label}</strong>.</p>
        {f'<p style="background:#f1f5f9;padding:10px;border-radius:6px;">{comentario}</p>' if comentario else ''}
        <p>Ingresa a la plataforma SGI Vanti para ver el detalle.</p>
    """)
    for u in recipientes:
        try:
            notify_user(u, subject, html, mensaje, tarea.id)
        except Exception as e:
            print(f"[ejecutivo.cambiar_estado] error notificando {u.username}: {e}")

    flash(f"Orden {tarea.orden} actualizada a '{estado_label}'.", 'success')
    return redirect(url_for('ejecutivo.dashboard'))


@ejecutivo_bp.route('/carta/<int:id>', methods=['GET', 'POST'])
@login_required
@role_required('ejecutivo')
def gestionar_carta(id):
    tarea = Imposibilidad.query.get_or_404(id)
    if tarea.ejecutivo_asignado != current_user.username:
        abort(403)
    if request.method == 'POST':
        guardar_datos_carta(tarea.carta, request.form)
        db.session.commit()
        flash('Datos de la carta actualizados.', 'success')
        return redirect(url_for('ejecutivo.dashboard'))
    return render_template('gestionar_carta.html', tarea=tarea)


@ejecutivo_bp.route('/carta/marcar_enviada/<int:id>', methods=['POST'])
@login_required
@role_required('ejecutivo')
def marcar_carta_enviada(id):
    tarea = Imposibilidad.query.get_or_404(id)
    if tarea.ejecutivo_asignado != current_user.username:
        abort(403)
    tarea.estado_tarea = 'carta_enviada'
    tarea.fecha_envio_carta = datetime.now()
    db.session.commit()

    # Notify firma
    firma_user = Usuario.query.filter_by(username=tarea.bp_firma).first()
    if firma_user:
        subject = f"Carta Enviada - Orden {tarea.orden}"
        html = render_template('email_base.html', content=f"""
            <p>Hola {firma_user.username},</p>
            <p>La carta para la orden <strong>{tarea.orden}</strong> ha sido marcada como <strong>Enviada</strong>.</p>
        """)
        wa_msg = f"Hola {firma_user.username}, la carta de la orden {tarea.orden} fue enviada."
        notify_user(firma_user, subject, html, wa_msg, tarea.id)

    flash('Carta marcada como "Enviada".', 'success')
    return redirect(url_for('ejecutivo.dashboard'))


@ejecutivo_bp.route('/carta/descargar_word/<int:id>')
@login_required
@role_required('ejecutivo')
def descargar_carta_word(id):
    tarea = Imposibilidad.query.get_or_404(id)
    if tarea.ejecutivo_asignado != current_user.username or not tarea.carta:
        abort(403)
    carta = tarea.carta
    document = Document()
    document.add_heading(f'Carta de Imposibilidad - Contrato: {tarea.cuenta_contrato}', level=1)
    document.add_paragraph(f"Bogotá, {datetime.now().strftime('%d de %B de %Y')}\n")
    document.add_paragraph(
        f"Señor(a):\n{carta.nombre_cliente or 'N/A'}\n"
        f"C.C. {carta.cedula_cliente or 'N/A'} de {carta.lugar_expedicion or 'N/A'}\n"
    )
    p = document.add_paragraph('Respetado(a) cliente,\n\n')
    p.add_run('Nos permitimos informarle sobre el estado de su solicitud...').bold = True

    if carta.observaciones_puntuales:
        document.add_paragraph(f"\nObservaciones: {carta.observaciones_puntuales}")
    if carta.direccion_predio:
        document.add_paragraph(f"Dirección del predio: {carta.direccion_predio}")
    if carta.distancia_acometida:
        document.add_paragraph(f"Distancia de acometida: {carta.distancia_acometida} m")
    if carta.tipo_avenida:
        document.add_paragraph(f"Tipo de vía: {carta.tipo_avenida}")

    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True,
                     download_name=f'carta_{tarea.cuenta_contrato}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@ejecutivo_bp.route('/descargar_cartas_zip')
@login_required
@role_required('ejecutivo')
def descargar_cartas_zip():
    tareas = Imposibilidad.query.filter_by(
        ejecutivo_asignado=current_user.username, tipo_tarea='carta'
    ).filter(Imposibilidad.estado_tarea != 'carta_enviada').all()

    if not tareas:
        flash("No hay cartas activas para descargar.", "info")
        return redirect(url_for('ejecutivo.dashboard'))

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for tarea in tareas:
            if not tarea.carta:
                continue
            doc_buffer = io.BytesIO()
            document = Document()
            document.add_heading(f'Carta - Contrato: {tarea.cuenta_contrato}', 1)
            document.add_paragraph(f"Cliente: {tarea.carta.nombre_cliente or 'N/A'}")
            document.add_paragraph(f"C.C.: {tarea.carta.cedula_cliente or 'N/A'}")
            document.add_paragraph(f"Dirección: {tarea.carta.direccion_predio or 'N/A'}")
            document.save(doc_buffer)
            doc_buffer.seek(0)
            zip_file.writestr(f"carta_{tarea.cuenta_contrato}.docx", doc_buffer.read())

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True,
                     download_name='cartas_pendientes.zip', mimetype='application/zip')
